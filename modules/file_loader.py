from langchain.schema import Document
import pandas as pd
import fitz  # PyMuPDF
import docx
from typing import List
import os
from utils.tools import _format_table
from utils.logger import Logger
import time
from config import get_config


# This is to suppress the warning from openpyxl
import warnings
warnings.filterwarnings("ignore", category=UserWarning, module='openpyxl')

class EnhancedDocumentLoader:
    def __init__(self):
        self.config = get_config()
        self.logger = Logger(log_file=self.config["loader"]["log_file"])
        self.supported_extensions = {
            '.pdf': self._load_pdf,
            '.docx': self._load_docx,
            '.csv': self._load_csv,
            '.xlsx': self._load_excel,
            '.xls': self._load_excel,
            '.xlsm': self._load_excel
        }
        
    def load_multiple_documents(self, files_path: List[str]) -> List[Document]:
        all_documents = []

        self.logger.info(f"Processing {len(files_path)} documents")
        start_time = time.time()

        for file_path in files_path:
            try:
                documents = self.load_document(file_path)
                all_documents.extend(documents)
                self.logger.info(f"✅ file loaded ({file_path})")
            except Exception as e:
                self.logger.info(f"❌ file failed loading ({file_path}) - ({str(e)})")        

        self.logger.metrics(
                "Document Processing", 
                start_time, 
                time.time(), 
                {"files_processed": len(files_path)}
            )
        return all_documents
            
    def load_document(self, file_path: str) -> List[Document]:
        """Main method to load any supported document type"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {ext}")
            
        return self.supported_extensions[ext](file_path)
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """Enhanced PDF loading with table detection"""
        doc = fitz.open(file_path)
        
        # Create a list to store documents, one per page
        documents = []
        
        for page_num, page in enumerate(doc):
            # Extract regular text
            text_content = page.get_text("text")
            
            # Extract tables
            tables = []
            try:
                # Get tables from the page
                tables_dict = page.find_tables()
                if tables_dict:
                    # tables_dict.tables contains the actual table objects
                    for table in tables_dict.tables:
                        # Extract table data as a list of lists
                        table_data = table.extract()
                        if table_data and len(table_data) > 1:  # Ensure table has content
                            # Convert table to structured format
                            table_text = _format_table(table_data)
                            tables.append(table_text)
            except AttributeError:
                # Handle older versions of PyMuPDF
                try:
                    rect_tables = page.find_tables()
                    for table in rect_tables:
                        table_data = table.extract()
                        if table_data and len(table_data) > 1:
                            table_text = _format_table(table_data)
                            tables.append(table_text)
                except Exception as e:
                    self.logger.warning(f"Warning: Table extraction failed on page {page_num + 1}: {e}")
            except Exception as e:
                self.logger.warning(f"Warning: Table extraction failed on page {page_num + 1}: {e}")
            
            # Create content for this page
            page_content = f"PAGE {page_num + 1}\n\n"
            if text_content.strip():
                page_content += f"\n{text_content}\n"
            if tables:
                page_content += "\n" + "\n\n".join(tables) + "\n\n"
            
            # Create metadata for this page
            metadata = {
                "source": file_path,
                "total_pages": len(doc),
                "page_number": page_num + 1
            }
            
            # Add document for this page to our list
            documents.append(Document(
                page_content=page_content,
                metadata=metadata
            ))
        
        doc.close()
        
        return documents
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """Parse Word document using python-docx with improved table handling and page-by-page splitting"""
        try:
            doc = docx.Document(file_path)
            documents = []
            current_page_content = []
            page_num = 1
            has_tables = False
            
            # Estimate page breaks based on paragraphs (approx. 3000 chars per page)
            chars_per_page = 3000
            char_count = 0
            
            # Process all paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                char_count += len(text)
                current_page_content.append(text)
                
                # Create a new page document when we exceed the character limit
                if char_count >= chars_per_page:
                    # Create metadata for this page
                    metadata = {
                        "source": file_path,
                        "has_tables": has_tables,
                        "content_type": "document",
                        "page_number": page_num
                    }
                    
                    # Create document for this page
                    documents.append(Document(
                        page_content=f"PAGE {page_num}\n\n" + "\n".join(current_page_content),
                        metadata=metadata
                    ))
                    
                    # Reset for next page
                    current_page_content = []
                    char_count = 0
                    page_num += 1
            
            # Process all tables with semantic formatting
            for table in doc.tables:
                has_tables = True
                table_content = []
                headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
                
                # Add table description
                table_content.append(f"\nThe table contains {len(table.rows)-1 if table.rows else 0} rows with columns: {', '.join(headers) if headers else 'no headers'}.")
                
                # Process each row
                for row_idx, row in enumerate(table.rows[1:] if table.rows else []):
                    row_parts = []
                    for cell_idx, cell in enumerate(row.cells):
                        header = headers[cell_idx] if cell_idx < len(headers) else f"Column {cell_idx+1}"
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_parts.append(f"{header} is {cell_text}")
                    
                    if row_parts:  # Only add non-empty rows
                        table_content.append(f"Row {row_idx+1} shows that " + ", ".join(row_parts) + ".")
                
                # Add table to a new page
                metadata = {
                    "source": file_path,
                    "has_tables": True,
                    "content_type": "document",
                    "page_number": page_num
                }
                
                documents.append(Document(
                    page_content=f"PAGE {page_num}\n\n" + "\n".join(table_content),
                    metadata=metadata
                ))
                page_num += 1
            
            # Add any remaining content as the final page
            if current_page_content:
                metadata = {
                    "source": file_path,
                    "has_tables": has_tables,
                    "content_type": "document",
                    "page_number": page_num,
                    "total_pages": page_num
                }
                
                documents.append(Document(
                    page_content=f"PAGE {page_num}\n\n" + "\n".join(current_page_content),
                    metadata=metadata
                ))
            
            # Update total_pages in all documents' metadata
            total_pages = len(documents)
            for doc in documents:
                doc.metadata["total_pages"] = total_pages
            
            return documents if documents else [Document(
                page_content="[Empty document]",
                metadata={"source": file_path, "total_pages": 1, "page_number": 1}
            )]
            
        except Exception as e:
            return [Document(
                page_content=f"Error loading DOCX: {str(e)}",
                metadata={"source": file_path, "error": True, "page_number": 1, "total_pages": 1}
            )]
    
    def _load_csv(self, file_path: str) -> List[Document]:
        """Enhanced CSV loading with row-by-row document creation"""
        df = pd.read_csv(file_path)
        documents = []
        
        # Get column headers
        headers = list(df.columns)
        
        # Process each row as a separate document
        for idx, row in df.iterrows():
            # Format row content
            row_content = f"ROW {idx + 1}:\n"
            for col in headers:
                value = row[col]
                if pd.notna(value):
                    row_content += f"{col}: {value}\n"
                else:
                    row_content += f"{col}: N/A\n"
            
            # Create metadata for this row
            metadata = {
                "source": file_path,
                "row_number": idx + 1,
                "total_rows": len(df),
                "columns": headers,
                "document_type": "csv_row"
            }
            
            # Create document for this row
            documents.append(Document(
                page_content=row_content,
                metadata=metadata
            ))
        
        return documents
    
    def _load_excel(self, file_path: str) -> List[Document]:
        """Load Excel file and create one document per row with human-readable content"""
        documents = []
        
        # Read all sheets from Excel file
        xlsx = pd.read_excel(file_path, sheet_name=None)
                    
        # Process each sheet
        for sheet_name, df in xlsx.items():
            if df.empty:
                self.logger.warning(f"Sheet '{sheet_name}' is empty")
                continue
            
            # Clean the dataframe by removing empty rows and columns
            df = df.dropna(how='all').dropna(axis=1, how='all')
            
            # Get column headers
            headers = list(df.columns)
            
            # Process each row
            for idx, row in df.iterrows():
                # Skip completely empty rows
                if row.isnull().all():
                    continue
                
                # Create human-readable content for this row
                content_parts = [
                    f"Sheet: {sheet_name}",
                    f"Row Number: {idx + 1}",
                    "\nRow Contents:"
                ]
                
                # Add each non-null cell value in a readable format
                for col in headers:
                    value = row[col]
                    if pd.notna(value):
                        content_parts.append(f"- {col}: {value}")
                
                # Join all parts with newlines
                row_content = "\n".join(content_parts)
                
                # Create metadata for this row
                metadata = {
                    "source": file_path,
                    "sheet_name": sheet_name,
                    "row_number": idx + 1,
                    "total_rows": len(df),
                    "columns": headers
                }
                
                # Create document for this row
                documents.append(Document(
                    page_content=row_content,
                    metadata=metadata
                ))
        
        
        return documents
    